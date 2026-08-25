import io
import os
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import List

from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import StreamingResponse

from lib.auth import current_user, log_activity, require_order
from lib.db import db
from models.schemas import (GenerateRequest, PackageRequest, ReportSave,
                            ReportSection, ReportVersion)

router = APIRouter(tags=["reports"])

NA = "Information not provided"
NA_DOCS = "Not available from the supplied documents"


def _v(value: str, fallback: str = NA) -> str:
    value = (value or "").strip()
    return value if value else fallback


async def _settings() -> dict:
    s = await db.settings.find_one({"id": "app"})
    return s or {"company_name": "Jed Red Solutions Pvt Ltd"}


def _kv(pairs: list) -> str:
    return "\n".join(f"{label}: {_v(value)}" for label, value in pairs)


async def _build_sections(order: dict) -> List[ReportSection]:
    docs = await db.documents.find({"order_id": order["id"], "deleted": {"$ne": True}}).to_list(500)
    docs.sort(key=lambda d: (d.get("timeline_index", 0), d.get("doc_date", "")))
    findings = await db.findings.find({"order_id": order["id"]}).to_list(500)
    assignee = await db.users.find_one({"id": order.get("assigned_to")}) if order.get("assigned_to") else None

    def info(d: dict, key: str) -> str:
        return (d.get("info") or {}).get(key, "") or ""

    docs_reviewed = "\n\n".join(
        f"{i + 1}. {_v(d.get('doc_type'), 'Document type not provided')} — Document No. {_v(d.get('doc_number'))}\n"
        f"   Document date: {_v(d.get('doc_date'))}\n"
        f"   Registration No. {_v(d.get('registration_number'))} at {_v(d.get('registration_office'))}\n"
        f"   Source: {_v(d.get('source'))} {('(' + d['source_url'] + ')') if d.get('source_url') else ''}\n"
        f"   Raw file attached: {_v(d.get('file_name'), 'No file attached')}"
        for i, d in enumerate(docs)) or NA_DOCS

    chain = "\n".join(
        f"{i + 1}. {_v(d.get('doc_date'))} — {_v(d.get('doc_type'), 'Document')} "
        f"(Doc No. {_v(d.get('doc_number'))}, Reg. No. {_v(d.get('registration_number'))}) — "
        f"Parties: {_v((info(d, 'seller') or info(d, 'donor') or info(d, 'mortgagor')), 'Party not provided')}"
        f" to {_v((info(d, 'buyer') or info(d, 'donee') or info(d, 'mortgagee')), 'Party not provided')}"
        f" — Property: {_v(info(d, 'prop_address') or info(d, 'survey_number'))}"
        for i, d in enumerate(docs)) or NA_DOCS

    ownership = "\n".join(
        f"- {_v(info(d, 'transaction_type') or d.get('doc_type'), 'Transaction')} dated "
        f"{_v(info(d, 'execution_date') or d.get('doc_date'))}: "
        f"{_v(info(d, 'seller') or info(d, 'donor'), 'Transferor not provided')} → "
        f"{_v(info(d, 'buyer') or info(d, 'donee'), 'Transferee not provided')}; "
        f"consideration {_v(info(d, 'consideration_amount'))}; extent {_v(info(d, 'extent_area'))}"
        for d in docs) or NA_DOCS

    important = "\n\n".join(
        f"[{_v(d.get('doc_type'), 'Document')} / {_v(d.get('doc_number'))}]\n{info(d, 'important_info').strip()}"
        for d in docs if info(d, 'important_info').strip()) or NA_DOCS

    findings_txt = "\n\n".join(
        f"- {_v(f.get('finding_type'), 'Finding')} ({_v(f.get('importance'), 'Importance not provided')}), "
        f"found {_v(f.get('date_found'))}\n  {_v(f.get('description'))}\n"
        f"  Source: {_v(f.get('source'))} {('(' + f['source_url'] + ')') if f.get('source_url') else ''}\n"
        f"  Researcher notes: {_v(f.get('researcher_notes'))}"
        for f in findings) or "No search findings were recorded for this order."

    issues = "\n".join(
        f"- [{_v(d.get('doc_type'), 'Document')} {_v(d.get('doc_number'))}] {info(d, 'potential_issues').strip()}"
        for d in docs if info(d, 'potential_issues').strip()) or "No potential issues or exceptions were recorded."

    notes = "\n\n".join(
        f"[{_v(d.get('doc_type'), 'Document')}] {info(d, 'researcher_notes').strip()}"
        for d in docs if info(d, 'researcher_notes').strip())
    observations = "\n\n".join(filter(None, [
        _v(order.get("additional_search_info"), "") and f"Additional search information: {order['additional_search_info']}",
        _v(order.get("internal_notes"), "") and f"Internal notes: {order['internal_notes']}",
        notes,
    ])) or NA

    sources = "\n".join(
        f"{i + 1}. {_v(d.get('file_name'), 'No file attached')} — {_v(d.get('doc_type'))}, "
        f"source: {_v(d.get('source'))} {('URL: ' + d['source_url']) if d.get('source_url') else ''}"
        for i, d in enumerate(docs)) or NA_DOCS

    return [
        ReportSection(heading="1. Order Information", body=_kv([
            ("Order ID", order.get("order_number")), ("Client name", order.get("client_name")),
            ("Client reference number", order.get("client_reference")),
            ("Order date", order.get("order_date")), ("Due date", order.get("due_date")),
            ("Priority", order.get("priority")), ("Order status", order.get("status")),
            ("Assigned researcher", assignee["name"] if assignee else ""),
        ])),
        ReportSection(heading="2. Property Details", body=_kv([
            ("Property owner name", order.get("property_owner")),
            ("Applicant / party name", order.get("applicant_name")),
            ("Property address", order.get("property_address")),
            ("Village / Town", order.get("village")), ("Taluk / Tehsil", order.get("taluk")),
            ("District", order.get("district")), ("State", order.get("state")),
            ("Survey number", order.get("survey_number")),
            ("Sub-division number", order.get("sub_division_number")),
            ("Plot number", order.get("plot_number")),
            ("Property ID / Khata number", order.get("khata_number")),
            ("Other identifying information", order.get("other_identifying_info")),
        ])),
        ReportSection(heading="3. Search Details", body=_kv([
            ("Registration district", order.get("registration_district")),
            ("Registration office", order.get("registration_office")),
            ("Search period", order.get("search_period")),
            ("Client order instructions / clues", order.get("client_instructions")),
        ]) + f"\n\nDocuments examined: {len(docs)}\nSearch findings recorded: {len(findings)}"),
        ReportSection(heading="4. Documents Reviewed", body=docs_reviewed),
        ReportSection(heading="5. Chronological Document History", body=chain),
        ReportSection(heading="6. Ownership / Transaction History", body=ownership),
        ReportSection(heading="7. Important Findings", body=important + "\n\n" + findings_txt),
        ReportSection(heading="8. Search Observations", body=observations),
        ReportSection(heading="9. Potential Issues / Exceptions", body=issues),
        ReportSection(heading="10. Conclusion / Summary", body=(
            f"This report summarises the documents and information entered into the system for order "
            f"{_v(order.get('order_number'))} in respect of the property described above. "
            f"{len(docs)} document(s) and {len(findings)} search finding(s) were recorded. "
            "Any field marked 'Information not provided' was not supplied and has not been assumed. "
            "No legal opinion or conclusion is expressed unless entered by an authorised user.")),
        ReportSection(heading="11. Source Documents", body=sources),
        ReportSection(heading="12. Researcher Information", body=_kv([
            ("Prepared by", assignee["name"] if assignee else ""),
            ("Researcher notes", notes),
        ])),
        ReportSection(heading="13. Reviewer / Approval Information", body=_kv([
            ("Reviewed by", ""), ("Approval status", order.get("status")), ("Approval date", ""),
        ])),
    ]


async def _ai_polish(order: dict, sections: List[ReportSection]) -> tuple[List[ReportSection], bool]:
    """Ask Gemini to tighten narrative sections only. Facts stay as entered."""
    key = os.environ.get("EMERGENT_LLM_KEY")
    if not key:
        return sections, False
    targets = {"10. Conclusion / Summary", "8. Search Observations", "7. Important Findings"}
    payload = "\n\n".join(f"### {s.heading}\n{s.body}" for s in sections if s.heading in targets)
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        chat = LlmChat(
            api_key=key, session_id=f"report-{order['id']}",
            system_message=(
                "You are a title-search report editor. Rewrite ONLY for clarity and professional tone. "
                "NEVER invent, infer or add facts, names, dates or numbers. Keep every phrase "
                "'Information not provided' / 'Not available from the supplied documents' exactly as-is. "
                "Return the same '### <heading>' blocks, nothing else, no preamble."),
        ).with_model("gemini", "gemini-3-flash-preview")
        reply = await chat.send_message(UserMessage(text=payload))
        text = reply if isinstance(reply, str) else str(reply)
        out: dict[str, str] = {}
        current = None
        for line in text.splitlines():
            if line.startswith("### "):
                current = line[4:].strip()
                out[current] = ""
            elif current:
                out[current] += line + "\n"
        polished = [ReportSection(heading=s.heading,
                                  body=(out.get(s.heading, "").strip() or s.body))
                    for s in sections]
        return polished, True
    except Exception:
        return sections, False


@router.get("/orders/{order_id}/report/versions", response_model=List[ReportVersion])
async def list_versions(order_id: str, request: Request):
    user = await current_user(request)
    await require_order(order_id, user, "view")
    rows = await db.report_versions.find({"order_id": order_id}).sort("version", -1).to_list(200)
    return [ReportVersion(**{k: v for k, v in r.items() if k != "_id"}) for r in rows]


async def _next_version(order_id: str) -> int:
    return await db.report_versions.count_documents({"order_id": order_id}) + 1


@router.post("/orders/{order_id}/report/generate", response_model=ReportVersion)
async def generate_report(order_id: str, payload: GenerateRequest, request: Request):
    user = await current_user(request)
    order = await require_order(order_id, user, "edit")
    sections = await _build_sections(order)
    ai_used = False
    if payload.use_ai:
        sections, ai_used = await _ai_polish(order, sections)
    version = ReportVersion(
        order_id=order_id, version=await _next_version(order_id),
        status="Draft", ai_generated=ai_used, sections=sections,
        created_by_name=user["name"], created_at=datetime.now(timezone.utc),
        change_note="AI-assisted draft generated" if ai_used else "Draft generated from entered data",
    )
    await db.report_versions.insert_one(version.model_dump())
    await db.orders.update_one({"id": order_id}, {"$set": {"status": "REPORT PREPARATION"}}) \
        if order["status"] in ("NEW", "IN PROGRESS", "DOCUMENTS COLLECTED") else None
    await log_activity(order_id, user, "Report generated", f"Version {version.version}")
    return version


@router.post("/orders/{order_id}/report/versions", response_model=ReportVersion)
async def save_version(order_id: str, payload: ReportSave, request: Request):
    user = await current_user(request)
    await require_order(order_id, user, "edit")
    version = ReportVersion(
        order_id=order_id, version=await _next_version(order_id), status="Revised",
        sections=payload.sections, created_by_name=user["name"],
        created_at=datetime.now(timezone.utc), change_note=payload.change_note or "Revised",
    )
    await db.report_versions.insert_one(version.model_dump())
    await log_activity(order_id, user, "Report edited", f"Version {version.version} saved")
    return version


@router.post("/orders/{order_id}/report/versions/{version_id}/approve", response_model=ReportVersion)
async def approve_version(order_id: str, version_id: str, request: Request):
    user = await current_user(request)
    await require_order(order_id, user, "approve")
    row = await db.report_versions.find_one({"id": version_id, "order_id": order_id})
    if not row:
        raise HTTPException(status_code=404, detail="Report version not found")
    sections = [ReportSection(**s) for s in row["sections"]]
    for s in sections:
        if s.heading.startswith("13."):
            s.body = (f"Reviewed by: {user['name']}\nApproval status: Approved\n"
                      f"Approval date: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}")
    approved = ReportVersion(
        order_id=order_id, version=await _next_version(order_id), status="Approved",
        sections=sections, created_by_name=user["name"], created_at=datetime.now(timezone.utc),
        change_note=f"Approved (from version {row['version']})",
    )
    await db.report_versions.insert_one(approved.model_dump())
    await db.orders.update_one({"id": order_id}, {"$set": {"status": "APPROVED"}})
    await log_activity(order_id, user, "Report approved", f"Version {approved.version}")
    return approved


async def _latest(order_id: str, approved_only: bool = False) -> dict:
    query: dict = {"order_id": order_id}
    if approved_only:
        query["status"] = "Approved"
    rows = await db.report_versions.find(query).sort("version", -1).to_list(1)
    if not rows:
        raise HTTPException(status_code=404, detail="No report has been generated for this order yet")
    return rows[0]


def _pdf_bytes(order: dict, report: dict, settings: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (PageBreak, Paragraph, SimpleDocTemplate,
                                    Spacer)

    buf = io.BytesIO()
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=9.5, leading=14)
    h = ParagraphStyle("h", parent=styles["Heading2"], fontSize=12, spaceBefore=10, spaceAfter=4)
    title = ParagraphStyle("t", parent=styles["Title"], fontSize=17)

    def footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.drawString(18 * mm, 12 * mm, f"{settings.get('company_name', '')} — {order.get('order_number', '')}")
        canvas.drawRightString(192 * mm, 12 * mm, f"Page {doc_.page}")
        canvas.restoreState()

    story = [
        Paragraph(settings.get("company_name", "Jed Red Solutions Pvt Ltd"), styles["Heading3"]),
        Paragraph("TITLE SEARCH SUMMARY REPORT", title),
        Spacer(1, 6),
        Paragraph(f"Order ID: <b>{order.get('order_number', '')}</b> &nbsp;|&nbsp; Client: <b>{order.get('client_name', '')}</b>", body),
        Paragraph(f"Property: {order.get('property_address') or NA}", body),
        Paragraph(f"Report date: {datetime.now(timezone.utc).strftime('%Y-%m-%d')} &nbsp;|&nbsp; "
                  f"Report version: {report.get('version')} ({report.get('status')})", body),
        Paragraph(f"Prepared by: {report.get('created_by_name') or NA}", body),
        Spacer(1, 10),
    ]
    for s in report.get("sections", []):
        story.append(Paragraph(s["heading"], h))
        for para in (s["body"] or NA).split("\n"):
            story.append(Paragraph(para.replace("&", "&amp;").replace("<", "&lt;") or "&nbsp;", body))
    if settings.get("report_footer"):
        story += [PageBreak(), Paragraph(settings["report_footer"], body)]
    SimpleDocTemplate(buf, pagesize=A4, topMargin=18 * mm, bottomMargin=20 * mm,
                      leftMargin=18 * mm, rightMargin=18 * mm,
                      title=f"Summary Report {order.get('order_number')}").build(
        story, onFirstPage=footer, onLaterPages=footer)
    return buf.getvalue()


def _docx_bytes(order: dict, report: dict, settings: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt

    d = Document()
    d.add_paragraph(settings.get("company_name", "Jed Red Solutions Pvt Ltd"))
    d.add_heading("TITLE SEARCH SUMMARY REPORT", level=0)
    meta = d.add_table(rows=0, cols=2)
    meta.style = "Table Grid"
    for label, value in [("Order ID", order.get("order_number", "")),
                         ("Client", order.get("client_name", "")),
                         ("Property", order.get("property_address") or NA),
                         ("Report date", datetime.now(timezone.utc).strftime("%Y-%m-%d")),
                         ("Report version", f"{report.get('version')} ({report.get('status')})"),
                         ("Prepared by", report.get("created_by_name") or NA),
                         ("Reviewed by", report.get("created_by_name") if report.get("status") == "Approved" else NA)]:
        row = meta.add_row().cells
        row[0].text = label
        row[1].text = str(value)
    for s in report.get("sections", []):
        d.add_heading(s["heading"], level=1)
        for para in (s["body"] or NA).split("\n"):
            p = d.add_paragraph(para)
            p.runs and setattr(p.runs[0].font, "size", Pt(10))
    if settings.get("report_footer"):
        d.add_paragraph(settings["report_footer"])
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


@router.get("/orders/{order_id}/report/export")
async def export_report(order_id: str, request: Request, fmt: str = "pdf"):
    user = await current_user(request)
    order = await require_order(order_id, user, "export")
    approved_only = user["role"] == "CLIENT"
    report = await _latest(order_id, approved_only=approved_only)
    settings = await _settings()
    if fmt == "docx":
        data = _docx_bytes(order, report, settings)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        name = f"Summary_Report_{order['order_number']}.docx"
    elif fmt == "pdf":
        data = _pdf_bytes(order, report, settings)
        media = "application/pdf"
        name = f"Summary_Report_{order['order_number']}.pdf"
    else:
        raise HTTPException(status_code=422, detail="Format must be pdf or docx")
    await log_activity(order_id, user, "Report exported", name)
    return StreamingResponse(io.BytesIO(data), media_type=media,
                             headers={"Content-Disposition": f'attachment; filename="{name}"'})


@router.post("/orders/{order_id}/package")
async def export_package(order_id: str, payload: PackageRequest, request: Request):
    user = await current_user(request)
    order = await require_order(order_id, user, "export")
    report = await _latest(order_id, approved_only=user["role"] == "CLIENT")
    settings = await _settings()
    docs = await db.documents.find({"order_id": order_id, "deleted": {"$ne": True}}).to_list(500)
    if payload.document_ids:
        docs = [d for d in docs if d["id"] in payload.document_ids]
    folder = order["order_number"]
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(f"{folder}/Summary_Report.pdf", _pdf_bytes(order, report, settings))
        z.writestr(f"{folder}/Summary_Report.docx", _docx_bytes(order, report, settings))
        if payload.include_index:
            index = ["DOCUMENT INDEX", f"Order: {folder}", f"Client: {order.get('client_name', '')}", ""]
            for i, d in enumerate(docs, 1):
                index.append(f"{i:03d}. {d.get('doc_type')} | Doc No. {d.get('doc_number') or NA} | "
                             f"Date {d.get('doc_date') or NA} | File: {d.get('file_name')}")
            z.writestr(f"{folder}/Document_Index.txt", "\n".join(index))
        for i, d in enumerate(docs, 1):
            path = d.get("stored_path")
            if path and os.path.exists(path):
                z.writestr(f"{folder}/Raw_Documents/Document_{i:03d}{Path(path).suffix}",
                           Path(path).read_bytes())
    await log_activity(order_id, user, "Order package exported", f"{len(docs)} raw document(s)")
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/zip", headers={
        "Content-Disposition": f'attachment; filename="{folder}_package.zip"'})
